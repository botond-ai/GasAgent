"""
Document parsing utilities for extracting text from various file formats.
Supports PDF and DOCX files.
"""
import io
import logging
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents and extract text content."""
    
    @staticmethod
    def parse_pdf(content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            content: PDF file content as bytes
        
        Returns:
            Extracted text content
        """
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")
            return full_text
        
        except Exception as e:
            logger.error(f"PDF parsing error: {e}", exc_info=True)
            raise ValueError(f"Failed to parse PDF: {e}")
    
    @staticmethod
    def parse_docx(content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            content: DOCX file content as bytes
        
        Returns:
            Extracted text content
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")
            return full_text
        
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}", exc_info=True)
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    @staticmethod
    def parse_document(content: bytes, mime_type: str) -> str:
        """
        Parse document based on MIME type.
        
        Args:
            content: File content as bytes
            mime_type: MIME type of the file
        
        Returns:
            Extracted text content
        
        Raises:
            ValueError: If MIME type is not supported
        """
        mime_type_lower = mime_type.lower()
        
        if mime_type_lower == 'application/pdf':
            return DocumentParser.parse_pdf(content)
        
        elif mime_type_lower in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            return DocumentParser.parse_docx(content)
        
        else:
            raise ValueError(
                f"Unsupported MIME type: {mime_type}. "
                f"Supported types: PDF, DOCX"
            )
    
    @staticmethod
    def get_document_metadata(text: str) -> dict:
        """
        Extract basic metadata from parsed text.
        
        Args:
            text: Parsed document text
        
        Returns:
            Dictionary with metadata (word count, character count, etc.)
        """
        words = text.split()
        lines = text.split('\n')
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'line_count': len(lines),
            'paragraph_count': len([line for line in lines if line.strip()]),
        }
