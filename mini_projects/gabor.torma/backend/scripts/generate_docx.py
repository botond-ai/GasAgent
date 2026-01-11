from docx import Document
import os

def generate_docx():
    doc = Document()
    doc.add_heading('Project Delta Architecture', 0)

    doc.add_paragraph('Meeting Minutes')

    doc.add_heading('Discussion', level=1)
    doc.add_paragraph('Frank: Using Microservices seems unnecessary for the MVP.')
    doc.add_paragraph('Grace: But we need to scale later.')
    doc.add_paragraph('Frank: OK, let\'s separate the auth service at least.')

    doc.add_heading('Action Items', level=1)
    doc.add_paragraph('Grace to draft the architecture diagram.', style='List Bullet')
    
    output_path = "mock_data/sample.docx"
    doc.save(output_path)
    print(f"Generated {output_path}")

if __name__ == "__main__":
    # Ensure directory exists (it should, but just in case)
    os.makedirs("mock_data", exist_ok=True)
    generate_docx()
